#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
RAG文件存储管理模块

本模块提供知识库文件的持久化存储、管理和预览功能。

支持两种存储后端：
- local: 本地文件系统 backend/data/rag_files/{year}/{month}/{filename}
- gcs: Google Cloud Storage gs://{bucket}/rag_files/{year}/{month}/{filename}

通过环境变量 RAG_STORAGE_BACKEND 和 RAG_GCS_BUCKET 切换。

作者: Project3 Team
版本: 2.0 (GCS support)
"""

import io
import os
import re
import tempfile
from datetime import datetime
from typing import Optional, Tuple

import pytz
from pathlib import Path

# Default short preview length (stored in DB)
DEFAULT_PREVIEW_LENGTH = 500


def _is_gcs_backend() -> bool:
    """Check if GCS storage backend is enabled."""
    try:
        from config.settings import RAG_STORAGE_BACKEND, RAG_GCS_BUCKET
        return RAG_STORAGE_BACKEND == "gcs" and bool(RAG_GCS_BUCKET)
    except Exception:
        return False


def _get_gcs_client():
    """Lazy import GCS client."""
    from google.cloud import storage
    return storage.Client()


def _get_gcs_blob(relative_path: str):
    """Get GCS blob for a relative path."""
    from config.settings import RAG_GCS_BUCKET
    client = _get_gcs_client()
    bucket = client.bucket(RAG_GCS_BUCKET)
    blob_path = f"rag_files/{relative_path}".replace("\\", "/")
    return bucket.blob(blob_path)


def sanitize_filename(original_filename: str, max_base_length: int = 120) -> str:
    """
    Sanitize user-provided filename to prevent traversal and unsafe characters.
    """
    name = Path(original_filename or "").name.replace("\x00", "").strip()
    if not name:
        return "upload.bin"
    base, ext = os.path.splitext(name)
    safe_base = re.sub(r"[^A-Za-z0-9._-]", "_", base).strip("._")
    safe_base = safe_base[:max_base_length] or "upload"
    safe_ext = re.sub(r"[^A-Za-z0-9.]", "", ext)[:10]
    return f"{safe_base}{safe_ext}"


def get_rag_storage_path() -> str:
    """
    获取RAG文件存储根目录（仅 local 后端有效）

    Returns:
        str: RAG文件存储根目录的绝对路径
    """
    current_dir = os.path.dirname(os.path.abspath(__file__))
    backend_dir = os.path.dirname(os.path.dirname(current_dir))
    storage_path = os.path.join(backend_dir, 'data', 'rag_files')

    # 确保目录存在（仅 local 时）
    if not _is_gcs_backend():
        os.makedirs(storage_path, exist_ok=True)

    return storage_path


def save_rag_file(file_content: bytes, original_filename: str) -> Tuple[str, str]:
    """
    保存RAG文件到存储目录（local 或 GCS）

    Args:
        file_content: 文件二进制内容
        original_filename: 原始文件名

    Returns:
        Tuple[str, str]: (完整路径/标识, 相对路径)
        - local: (绝对路径, 相对路径如 2024/02/xxx.xlsx)
        - gcs: (gs://bucket/rag_files/..., 相对路径)
    """
    beijing_tz = pytz.timezone('Asia/Shanghai')
    now = datetime.now(beijing_tz)
    year = now.strftime('%Y')
    month = now.strftime('%m')
    timestamp = int(now.timestamp())
    safe_original_filename = sanitize_filename(original_filename)
    filename_parts = os.path.splitext(safe_original_filename)
    unique_filename = f"{filename_parts[0]}_{timestamp}{filename_parts[1]}"
    relative_path = os.path.join(year, month, unique_filename).replace("\\", "/")

    if _is_gcs_backend():
        blob = _get_gcs_blob(relative_path)
        blob.upload_from_string(
            file_content,
            content_type="application/octet-stream",
        )
        from config.settings import RAG_GCS_BUCKET
        full_path = f"gs://{RAG_GCS_BUCKET}/rag_files/{relative_path}"
        print(f"✅ 文件保存成功 (GCS): {relative_path}")
        return full_path, relative_path

    # Local backend
    storage_root = get_rag_storage_path()
    storage_dir = os.path.join(storage_root, year, month)
    os.makedirs(storage_dir, exist_ok=True)
    full_path = os.path.join(storage_dir, unique_filename)
    with open(full_path, 'wb') as f:
        f.write(file_content)
    print(f"✅ 文件保存成功: {relative_path}")
    return full_path, relative_path


def delete_rag_file(file_path: str) -> bool:
    """
    删除RAG文件

    Args:
        file_path: 相对路径（如 2024/02/xxx.xlsx）或绝对路径（local）

    Returns:
        bool: 删除成功返回True，失败返回False
    """
    try:
        if _is_gcs_backend():
            # file_path 应为相对路径
            rel = file_path
            if os.path.isabs(file_path):
                # 尝试从路径中提取相对部分
                if "rag_files" in file_path:
                    rel = file_path.split("rag_files/", 1)[-1].replace("\\", "/")
                else:
                    rel = os.path.basename(file_path)
            blob = _get_gcs_blob(rel)
            if blob.exists():
                blob.delete()
                print(f"✅ 文件删除成功 (GCS): {rel}")
                return True
            print(f"⚠️ 文件不存在 (GCS): {rel}")
            return False

        # Local backend
        if not os.path.isabs(file_path):
            storage_root = get_rag_storage_path()
            file_path = os.path.join(storage_root, file_path)

        if os.path.exists(file_path):
            os.remove(file_path)
            print(f"✅ 文件删除成功: {file_path}")
            _clean_empty_dirs(os.path.dirname(file_path))
            return True
        print(f"⚠️ 文件不存在: {file_path}")
        return False

    except Exception as e:
        print(f"❌ 文件删除失败: {e}")
        return False


def _clean_empty_dirs(directory: str):
    """清理空目录（仅 local 后端）"""
    try:
        if _is_gcs_backend():
            return
        storage_root = get_rag_storage_path()
        if not directory.startswith(storage_root):
            return
        if os.path.exists(directory) and not os.listdir(directory) and directory != storage_root:
            os.rmdir(directory)
            print(f"🗑️ 清理空目录: {directory}")
            _clean_empty_dirs(os.path.dirname(directory))
    except Exception as e:
        print(f"⚠️ 清理目录时出错: {e}")


def read_file_bytes(file_path: str) -> Optional[bytes]:
    """
    读取文件内容为 bytes。支持 local 和 GCS。

    Args:
        file_path: 相对路径或绝对路径（local）

    Returns:
        Optional[bytes]: 文件内容，失败返回 None
    """
    try:
        if _is_gcs_backend():
            rel = file_path
            if os.path.isabs(file_path) and "rag_files" in file_path:
                rel = file_path.split("rag_files/", 1)[-1].replace("\\", "/")
            elif os.path.isabs(file_path):
                rel = os.path.basename(file_path)
            blob = _get_gcs_blob(rel)
            if not blob.exists():
                return None
            return blob.download_as_bytes()
        # Local
        if not os.path.isabs(file_path):
            storage_root = get_rag_storage_path()
            file_path = os.path.join(storage_root, file_path)
        if not os.path.exists(file_path):
            return None
        with open(file_path, 'rb') as f:
            return f.read()
    except Exception as e:
        print(f"❌ 读取文件失败: {e}")
        return None


def get_local_path_for_reading(file_path: str) -> Optional[str]:
    """
    获取可用于读取的本地路径。GCS 时会下载到临时文件。

    调用方使用完毕后，临时文件由系统清理（进程退出时）。

    Args:
        file_path: 相对路径

    Returns:
        Optional[str]: 本地文件路径，失败返回 None
    """
    try:
        if _is_gcs_backend():
            content = read_file_bytes(file_path)
            if content is None:
                return None
            fd, path = tempfile.mkstemp(suffix=os.path.splitext(file_path)[1])
            try:
                os.write(fd, content)
                return path
            finally:
                os.close(fd)
        # Local
        return get_absolute_path(file_path)
    except Exception as e:
        print(f"❌ get_local_path_for_reading 失败: {e}")
        return None


def _preview_from_bytes(content: bytes, file_type: str, max_length: int) -> Optional[str]:
    """从 bytes 生成预览文本（供 GCS 使用）"""
    try:
        preview_text = ""
        if file_type in ['txt', 'csv']:
            preview_text = content.decode('utf-8', errors='ignore')[:max_length]
        elif file_type == 'excel':
            try:
                import pandas as pd
                df = pd.read_excel(io.BytesIO(content), sheet_name=0, nrows=10)
                preview_text = df.to_string()[:max_length]
            except Exception as e:
                print(f"⚠️ Excel预览失败: {e}")
                preview_text = "[Excel文件，无法生成预览]"
        elif file_type == 'word':
            try:
                from docx import Document
                doc = Document(io.BytesIO(content))
                paragraphs = [para.text for para in doc.paragraphs[:5]]
                preview_text = "\n".join(paragraphs)[:max_length]
            except Exception as e:
                print(f"⚠️ Word预览失败: {e}")
                preview_text = "[Word文件，无法生成预览]"
        elif file_type == 'pdf':
            try:
                import pdfplumber
                with pdfplumber.open(io.BytesIO(content)) as pdf:
                    if len(pdf.pages) > 0:
                        text = pdf.pages[0].extract_text()
                        preview_text = text[:max_length] if text else "[PDF文件，无文本内容]"
                    else:
                        preview_text = "[PDF文件为空]"
            except Exception as e:
                print(f"⚠️ PDF预览失败: {e}")
                preview_text = "[PDF文件，无法生成预览]"
        elif file_type == 'image':
            preview_text = "[图片文件]"
        else:
            preview_text = f"[{file_type}文件，不支持预览]"
        return preview_text
    except Exception as e:
        print(f"❌ _preview_from_bytes 失败: {e}")
        return None


def get_file_preview(file_path: str, file_type: str, max_length: int = 500) -> Optional[str]:
    """
    生成文件预览内容

    Args:
        file_path: 相对路径或绝对路径
        file_type: 文件类型（excel, word, pdf, txt, csv等）
        max_length: 最大预览字符数（默认500）

    Returns:
        Optional[str]: 预览文本，失败返回None
    """
    try:
        if _is_gcs_backend():
            content = read_file_bytes(file_path)
            if content is None:
                print(f"⚠️ 文件不存在 (GCS): {file_path}")
                return None
            return _preview_from_bytes(content, file_type, max_length)

        # Local
        if not os.path.isabs(file_path):
            storage_root = get_rag_storage_path()
            file_path = os.path.join(storage_root, file_path)
        if not os.path.exists(file_path):
            print(f"⚠️ 文件不存在: {file_path}")
            return None

        preview_text = ""
        if file_type in ['txt', 'csv']:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                preview_text = f.read(max_length)
        elif file_type == 'excel':
            try:
                import pandas as pd
                df = pd.read_excel(file_path, sheet_name=0, nrows=10)
                preview_text = df.to_string()[:max_length]
            except Exception as e:
                print(f"⚠️ Excel预览失败: {e}")
                preview_text = "[Excel文件，无法生成预览]"
        elif file_type == 'word':
            try:
                from docx import Document
                doc = Document(file_path)
                paragraphs = [para.text for para in doc.paragraphs[:5]]
                preview_text = "\n".join(paragraphs)[:max_length]
            except Exception as e:
                print(f"⚠️ Word预览失败: {e}")
                preview_text = "[Word文件，无法生成预览]"
        elif file_type == 'pdf':
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    if len(pdf.pages) > 0:
                        text = pdf.pages[0].extract_text()
                        preview_text = text[:max_length] if text else "[PDF文件，无文本内容]"
                    else:
                        preview_text = "[PDF文件为空]"
            except Exception as e:
                print(f"⚠️ PDF预览失败: {e}")
                preview_text = "[PDF文件，无法生成预览]"
        elif file_type == 'image':
            preview_text = "[图片文件]"
        else:
            preview_text = f"[{file_type}文件，不支持预览]"
        return preview_text
    except Exception as e:
        print(f"❌ 生成预览失败: {e}")
        return None


def _full_text_from_bytes(content: bytes, file_type: str) -> Tuple[Optional[str], int]:
    """从 bytes 提取全文（供 get_file_preview_slice GCS 使用）"""
    full_text = None
    if file_type in ['txt', 'csv']:
        full_text = content.decode('utf-8', errors='ignore')
    elif file_type == 'excel':
        try:
            import pandas as pd
            df = pd.read_excel(io.BytesIO(content), sheet_name=0)
            full_text = df.to_string()
        except Exception as e:
            print(f"⚠️ Excel 全量读取失败: {e}")
            return "[Excel文件，无法生成全文预览]", 0
    elif file_type == 'word':
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            full_text = "\n".join(p.text for p in doc.paragraphs)
        except Exception as e:
            print(f"⚠️ Word 全量读取失败: {e}")
            return "[Word文件，无法生成全文预览]", 0
    elif file_type == 'pdf':
        try:
            import pdfplumber
            parts = []
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        parts.append(t)
            full_text = "\n\n".join(parts) if parts else "[PDF文件，无文本内容]"
        except Exception as e:
            print(f"⚠️ PDF 全量读取失败: {e}")
            return "[PDF文件，无法生成全文预览]", 0
    elif file_type == 'image':
        return "[图片文件]", 0
    else:
        return f"[{file_type}文件，不支持预览]", 0
    if full_text is None:
        return None, 0
    return full_text, len(full_text)


def get_file_preview_slice(
    file_path: str,
    file_type: str,
    offset: int = 0,
    limit: Optional[int] = None,
) -> Tuple[Optional[str], int]:
    """
    获取文件预览内容（支持分页/全量）。支持 local 和 GCS。

    Args:
        file_path: 相对路径或绝对路径
        file_type: 文件类型
        offset: 起始字符偏移
        limit: 返回最大字符数，None 表示全量

    Returns:
        Tuple[Optional[str], int]: (预览文本, 文件总字符数)，失败时 (None, 0)
    """
    try:
        if _is_gcs_backend():
            content = read_file_bytes(file_path)
            if content is None:
                print(f"⚠️ 文件不存在 (GCS): {file_path}")
                return None, 0
            full_text, total = _full_text_from_bytes(content, file_type)
            if full_text is None:
                return None, total
            if total == 0:
                return full_text, 0
            if limit is None:
                snippet = full_text[offset:] if offset > 0 else full_text
            else:
                end = min(offset + limit, total)
                snippet = full_text[offset:end]
            return snippet, total

        # Local
        if not os.path.isabs(file_path):
            storage_root = get_rag_storage_path()
            file_path = os.path.join(storage_root, file_path)
        if not os.path.exists(file_path):
            print(f"⚠️ 文件不存在: {file_path}")
            return None, 0

        full_text = None
        if file_type in ['txt', 'csv']:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                full_text = f.read()
        elif file_type == 'excel':
            try:
                import pandas as pd
                df = pd.read_excel(file_path, sheet_name=0)
                full_text = df.to_string()
            except Exception as e:
                print(f"⚠️ Excel 全量读取失败: {e}")
                return "[Excel文件，无法生成全文预览]", 0
        elif file_type == 'word':
            try:
                from docx import Document
                doc = Document(file_path)
                full_text = "\n".join(p.text for p in doc.paragraphs)
            except Exception as e:
                print(f"⚠️ Word 全量读取失败: {e}")
                return "[Word文件，无法生成全文预览]", 0
        elif file_type == 'pdf':
            try:
                import pdfplumber
                parts = []
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        t = page.extract_text()
                        if t:
                            parts.append(t)
                full_text = "\n\n".join(parts) if parts else "[PDF文件，无文本内容]"
            except Exception as e:
                print(f"⚠️ PDF 全量读取失败: {e}")
                return "[PDF文件，无法生成全文预览]", 0
        elif file_type == 'image':
            return "[图片文件]", 0
        else:
            return f"[{file_type}文件，不支持预览]", 0

        if full_text is None:
            return None, 0
        total = len(full_text)
        if limit is None:
            snippet = full_text[offset:] if offset > 0 else full_text
        else:
            end = min(offset + limit, total)
            snippet = full_text[offset:end]
        return snippet, total

    except Exception as e:
        print(f"❌ get_file_preview_slice 失败: {e}")
        return None, 0


def get_absolute_path(relative_path: str) -> str:
    """
    将相对路径转换为可识别的完整路径。
    - local: 返回本地绝对路径
    - gcs: 返回 gs://bucket/rag_files/... 格式
    """
    if os.path.isabs(relative_path):
        return relative_path
    if _is_gcs_backend():
        from config.settings import RAG_GCS_BUCKET
        return f"gs://{RAG_GCS_BUCKET}/rag_files/{relative_path}".replace("\\", "/")
    storage_root = get_rag_storage_path()
    return os.path.join(storage_root, relative_path)


def file_exists(file_path: str) -> bool:
    """
    检查文件是否存在。支持 local 和 GCS。
    """
    try:
        if _is_gcs_backend():
            rel = file_path
            if os.path.isabs(file_path) and "rag_files" in file_path:
                rel = file_path.split("rag_files/", 1)[-1].replace("\\", "/")
            elif os.path.isabs(file_path):
                rel = os.path.basename(file_path)
            blob = _get_gcs_blob(rel)
            return blob.exists()
        if not os.path.isabs(file_path):
            file_path = get_absolute_path(file_path)
        return os.path.exists(file_path)
    except Exception:
        return False


def get_file_size(file_path: str) -> Optional[int]:
    """
    获取文件大小（bytes）。支持 local 和 GCS。
    """
    try:
        if _is_gcs_backend():
            rel = file_path
            if os.path.isabs(file_path) and "rag_files" in file_path:
                rel = file_path.split("rag_files/", 1)[-1].replace("\\", "/")
            elif os.path.isabs(file_path):
                rel = os.path.basename(file_path)
            blob = _get_gcs_blob(rel)
            if blob.exists():
                blob.reload()
                return blob.size
            return None
        if not os.path.isabs(file_path):
            file_path = get_absolute_path(file_path)
        if os.path.exists(file_path):
            return os.path.getsize(file_path)
        return None
    except Exception as e:
        print(f"❌ 获取文件大小失败: {e}")
        return None
